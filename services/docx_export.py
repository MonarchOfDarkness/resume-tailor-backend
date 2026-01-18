import io
from docx import Document
from docx.shared import Pt

def export_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    p = doc.add_paragraph((title or "").strip())
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    for line in (content or "").splitlines():
        doc.add_paragraph(line.rstrip() if line.strip() else "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
